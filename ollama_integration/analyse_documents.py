import time
import os

from database.create_db_connection import get_connection
from ollama import ChatResponse, chat, Client
from docx import Document
from pathlib import Path
from io import BytesIO

DIRECTORY_PATH = Path("storage/docx_downloads").resolve()
TABLE_NAME = "Acts"
TABLE_COLUMN = "ollamaAnalysedDocument"

def get_docx_files():
    files = []
    for filename in os.listdir(DIRECTORY_PATH):
        files.append(filename)
    print(f"All files: {len(files)}")

    return files

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])

    return text

def main():
    with open('ollama_integration/AI_vertinimas.txt', 'r') as file:
        evaluation_doc_lines = file.readlines()
    evluation_text = ''.join(evaluation_doc_lines)

    client = Client(host='http://localhost:11434') 
    files = get_docx_files()

    for file_name in files:
        try:
            base_dir = Path(__file__).parent.resolve().parent
            file_path = f'{base_dir}/storage/docx_downloads/{file_name}'

            # Pass file to Ollama
            act_content = extract_text_from_docx(file_path)
            prompt = f"""Išanalizuok pagal pateiktus kriterijus kiekvienai kategorijai, o paskui atsakyk į apačioje pateiktus klausimus: 
                        {evluation_text} įkeltam dokumentui: {act_content}. Atsakymas turi atsirasti prie kiekvieno parašyto vertinimo punkto."""
            print(prompt)

            start = time.time()
        
            response_doc = Document()

            response: ChatResponse = client.chat(model='llama3.2', messages=[ {'role': 'user', 'content': prompt} ])
            print(response['message']['content'])
            response_content = response['message']['content']

            # Add response into doc
            response_doc.add_paragraph(response_content)

            # Save the document to a BytesIO object (in memory)
            doc_buffer = BytesIO()
            response_doc.save(doc_buffer)
            doc_buffer.seek(0)  # Reset the pointer to the start of the buffer

            elapsed_time = time.time() - start
            print(f"Time elapsed: {int(elapsed_time // 60)} minutes {int(elapsed_time % 60)} seconds")

            # Add file into db
            connection = get_connection()
            cursor = connection.cursor()

            file_act_id = file_name.split("_")[-1].split(".")[0]

            # Insert analysed file into Acts table
            cursor.execute(f"UPDATE {TABLE_NAME} SET {TABLE_COLUMN} = ? WHERE act_id = ?", (doc_buffer.read(), file_act_id))
            connection.commit()
        except Exception as e:
            print(f"Error: {e}")

if __name__ == "__main__":
    main()